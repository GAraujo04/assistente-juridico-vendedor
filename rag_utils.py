"""
rag_utils.py
------------
Funções auxiliares de RAG (Retrieval-Augmented Generation) usadas pelo app.py:

- Extração de texto de PDF/DOCX
- Divisão do texto em chunks
- Índice vetorial local (ChromaDB) com embeddings locais (SentenceTransformers)
- Persistência em disco dos contratos (texto, arquivo original, resumo)
- Chamadas ao LLM (OpenAI) para resumo e para responder perguntas
"""

from __future__ import annotations

import hashlib
import io
import json
import time
from pathlib import Path

import chromadb
import docx
from pypdf import PdfReader

# ---------------------------------------------------------------------------
# 0) Armazenamento em disco dos contratos (persiste entre reinícios do app)
# ---------------------------------------------------------------------------

DATA_DIR = Path(__file__).parent / "data"
CONTRACTS_DIR = DATA_DIR / "contracts"
CHROMA_DIR = DATA_DIR / "chroma"

_chroma_client = None


def get_chroma_client():
    """Client único do ChromaDB, persistido em disco (data/chroma)."""
    global _chroma_client
    if _chroma_client is None:
        CHROMA_DIR.mkdir(parents=True, exist_ok=True)
        _chroma_client = chromadb.PersistentClient(path=str(CHROMA_DIR))
    return _chroma_client


def compute_contract_id(file_bytes: bytes) -> str:
    """ID estável baseado no conteúdo do arquivo (o mesmo contrato reenviado reaproveita os dados)."""
    return hashlib.sha256(file_bytes).hexdigest()[:16]


def save_contract(contract_id: str, filename: str, file_bytes: bytes, full_text: str) -> None:
    """Salva o arquivo original, o texto extraído e os metadados do contrato."""
    folder = CONTRACTS_DIR / contract_id
    folder.mkdir(parents=True, exist_ok=True)

    ext = Path(filename).suffix
    (folder / f"original{ext}").write_bytes(file_bytes)
    (folder / "text.txt").write_text(full_text, encoding="utf-8")

    meta_path = folder / "meta.json"
    if not meta_path.exists():
        meta = {"filename": filename, "uploaded_at": time.time()}
        meta_path.write_text(json.dumps(meta, ensure_ascii=False), encoding="utf-8")


def save_summary(contract_id: str, summary: str) -> None:
    (CONTRACTS_DIR / contract_id / "summary.md").write_text(summary, encoding="utf-8")


def load_summary(contract_id: str) -> str | None:
    path = CONTRACTS_DIR / contract_id / "summary.md"
    return path.read_text(encoding="utf-8") if path.exists() else None


def load_contract_text(contract_id: str) -> str:
    path = CONTRACTS_DIR / contract_id / "text.txt"
    return path.read_text(encoding="utf-8") if path.exists() else ""


def list_stored_contracts() -> list[dict]:
    """Lista os contratos já salvos em disco, mais recentes primeiro."""
    if not CONTRACTS_DIR.exists():
        return []

    items = []
    for folder in CONTRACTS_DIR.iterdir():
        meta_path = folder / "meta.json"
        if not meta_path.exists():
            continue
        meta = json.loads(meta_path.read_text(encoding="utf-8"))
        items.append(
            {
                "id": folder.name,
                "filename": meta.get("filename", folder.name),
                "uploaded_at": meta.get("uploaded_at", 0),
                "has_summary": (folder / "summary.md").exists(),
            }
        )

    items.sort(key=lambda item: item["uploaded_at"], reverse=True)
    return items


# ---------------------------------------------------------------------------
# 1) Extração de texto do contrato
# ---------------------------------------------------------------------------

def extract_text(uploaded_file) -> str:
    """Recebe um arquivo enviado via Streamlit (.pdf ou .docx) e retorna o texto puro."""
    filename = uploaded_file.name.lower()
    data = uploaded_file.read()

    if filename.endswith(".pdf"):
        return _extract_pdf(data)
    elif filename.endswith(".docx"):
        return _extract_docx(data)
    else:
        raise ValueError("Formato não suportado. Envie um arquivo .pdf ou .docx.")


def _extract_pdf(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(pages)


def _extract_docx(data: bytes) -> str:
    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]

    # Também aproveita o texto de tabelas (comum em contratos: valores, prazos etc.)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)

    return "\n".join(parts)


# ---------------------------------------------------------------------------
# 2) Divisão em chunks
# ---------------------------------------------------------------------------

def split_into_chunks(text: str, chunk_size: int = 800, overlap: int = 100) -> list[str]:
    """Divide o texto em pedaços menores (com sobreposição) para facilitar a busca semântica."""
    text = " ".join(text.split())  # normaliza espaços/quebras de linha
    if not text:
        return []

    chunks = []
    start = 0
    length = len(text)

    while start < length:
        end = min(start + chunk_size, length)
        chunks.append(text[start:end])
        if end == length:
            break
        start = end - overlap

    return chunks


# ---------------------------------------------------------------------------
# 3) Índice vetorial local (ChromaDB + SentenceTransformers)
# ---------------------------------------------------------------------------
#
# Todos os contratos ficam na MESMA coleção (não uma por contrato), com
# metadados indicando de qual contrato veio cada trecho. Isso permite que o
# chat pesquise e cruze informações entre TODOS os contratos já salvos, e não
# só o que está "aberto" na tela no momento.

ALL_CONTRACTS_COLLECTION = "contratos"


class ContractStore:
    """Índice vetorial único, compartilhado por todos os contratos salvos."""

    def __init__(self, embedding_model):
        self.embedding_model = embedding_model
        self._client = get_chroma_client()
        self._collection = self._client.get_or_create_collection(ALL_CONTRACTS_COLLECTION)

    def is_indexed(self, contract_id: str) -> bool:
        existing = self._collection.get(ids=[f"{contract_id}-chunk-0"])
        return len(existing.get("ids", [])) > 0

    def add_chunks(self, contract_id: str, filename: str, chunks: list[str]) -> None:
        if not chunks or self.is_indexed(contract_id):
            return  # já indexado (contrato reaberto/reenviado) ou nada para indexar
        embeddings = self.embedding_model.encode(chunks, show_progress_bar=False).tolist()
        ids = [f"{contract_id}-chunk-{i}" for i in range(len(chunks))]
        metadatas = [{"contract_id": contract_id, "filename": filename} for _ in chunks]
        self._collection.add(documents=chunks, embeddings=embeddings, ids=ids, metadatas=metadatas)

    def query(self, question: str, top_k: int = 5, contract_id: str | None = None) -> list[dict]:
        """Busca nos trechos de todos os contratos (ou só de um, se `contract_id` for passado)."""
        if self._collection.count() == 0:
            return []
        query_embedding = self.embedding_model.encode([question]).tolist()
        where = {"contract_id": contract_id} if contract_id else None
        results = self._collection.query(query_embeddings=query_embedding, n_results=top_k, where=where)
        documents = results.get("documents", [[]])[0]
        metadatas = results.get("metadatas", [[]])[0]
        return [
            {"text": doc, "filename": meta.get("filename", "contrato")}
            for doc, meta in zip(documents, metadatas)
        ]


def sync_all_contracts_to_store(store: ContractStore) -> None:
    """Garante que todo contrato salvo em disco esteja indexado na coleção compartilhada.

    Cobre o caso de contratos salvos antes dessa mudança (ou qualquer
    inconsistência entre `data/contracts` e o índice vetorial).
    """
    for item in list_stored_contracts():
        if store.is_indexed(item["id"]):
            continue
        full_text = load_contract_text(item["id"])
        chunks = split_into_chunks(full_text)
        store.add_chunks(item["id"], item["filename"], chunks)


# ---------------------------------------------------------------------------
# 4) Prompts e chamadas ao LLM
# ---------------------------------------------------------------------------

SUMMARY_PROMPT = """Você é um assistente jurídico especializado em apoiar VENDEDORES durante negociações de contratos.
Analise o contrato abaixo e produza um resumo objetivo em Markdown, com exatamente estas seções:

### 🤝 Partes Envolvidas
### 💰 Valor Total e Forma de Pagamento
### 📄 Cláusulas de Rescisão, Multas e Prazos
### ⚠️ Riscos e Pontos de Atenção para o Vendedor

Regras:
- Use bullet points curtos e linguagem simples (o leitor é um vendedor, não um advogado).
- Baseie-se ESTRITAMENTE no texto do contrato abaixo.
- Se alguma informação não estiver presente, escreva "Não especificado no contrato".

CONTRATO:
\"\"\"
{contract_text}
\"\"\"
"""

ANSWER_SYSTEM_PROMPT = """Você é um assistente jurídico que ajuda vendedores a entender contratos antes de fechar uma venda.

Você recebe trechos de um ou mais contratos já salvos como contexto (cada trecho vem marcado com "[Contrato: nome-do-arquivo]") e a pergunta do vendedor. Siga estas regras:

1. Se a pergunta for sobre o conteúdo de algum contrato (cláusulas, valores, prazos, partes envolvidas, riscos etc.), responda ESTRITAMENTE com base nos trechos fornecidos. Nunca invente cláusulas, valores, prazos ou condições que não estejam no contexto. Sempre que responder com base em um trecho, diga de qual contrato ele veio (pelo nome do arquivo). Se a informação perguntada não estiver no contexto, diga claramente: "Essa informação não foi encontrada nos contratos salvos."
2. Se os trechos vierem de mais de um contrato e não for óbvio qual deles o vendedor quer, liste rapidamente as opções encontradas (nome do arquivo + um dado que ajude a diferenciar) antes de aprofundar.
3. Se a pergunta for geral — dúvida jurídica genérica, de técnica de vendas, ou qualquer outro assunto que não dependa de um contrato específico — responda normalmente usando seu conhecimento geral. Nesse caso, deixe claro que a resposta não vem de nenhum contrato salvo.
4. Seja claro, direto e use linguagem acessível para um vendedor, não um advogado."""


def generate_summary(client, model: str, contract_text: str, max_chars: int = 18000) -> str:
    """Gera o resumo automático inicial do contrato completo (truncado se muito grande)."""
    truncated = contract_text[:max_chars]
    response = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": "Você é um assistente jurídico objetivo e cuidadoso."},
            {"role": "user", "content": SUMMARY_PROMPT.format(contract_text=truncated)},
        ],
        temperature=0.2,
    )
    return response.choices[0].message.content


def answer_question(client, model: str, question: str, context_chunks: list[dict]) -> str:
    """Responde a uma pergunta do vendedor usando trechos recuperados de um ou mais contratos salvos.

    `context_chunks` é uma lista de dicts {"text": ..., "filename": ...}, já
    ordenada por relevância (pode conter trechos de contratos diferentes).
    """
    if context_chunks:
        context = "\n\n---\n\n".join(
            f"[Contrato: {item['filename']}]\n{item['text']}" for item in context_chunks
        )
    else:
        context = "Nenhum contrato foi enviado, ou nenhum trecho relevante foi encontrado nos contratos salvos para esta pergunta."

    user_prompt = f"""Trechos relevantes dos contratos salvos:
\"\"\"
{context}
\"\"\"

Pergunta do vendedor: {question}
"""
    response = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": ANSWER_SYSTEM_PROMPT},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0.1,
    )
    return response.choices[0].message.content
